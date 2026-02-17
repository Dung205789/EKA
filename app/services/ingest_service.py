import uuid
from pathlib import Path
from urllib.parse import urlparse

from app.core.models import Document


def ingest_txt_path(path: str, *, title_override: str | None = None, meta_extra: dict | None = None) -> Document:
    p = Path(path)
    text = p.read_text(encoding="utf-8", errors="ignore")
    meta = {"path": str(p)}
    if meta_extra:
        meta.update(meta_extra)
    return Document(
        doc_id=str(uuid.uuid4()),
        source="txt",
        title=(title_override or p.name),
        raw_text=text,
        meta=meta,
    )


def ingest_docx_path(path: str, *, title_override: str | None = None, meta_extra: dict | None = None) -> Document:
    from docx import Document as Docx

    p = Path(path)
    d = Docx(path)
    parts = []
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    text = "\n".join(parts)
    meta = {"path": str(p)}
    if meta_extra:
        meta.update(meta_extra)
    return Document(
        doc_id=str(uuid.uuid4()),
        source="docx",
        title=(title_override or p.name),
        raw_text=text,
        meta=meta,
    )


def ingest_pdf_path(path: str, *, title_override: str | None = None, meta_extra: dict | None = None) -> Document:
    from pypdf import PdfReader

    p = Path(path)
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    text = "\n\n".join(parts)
    meta = {"path": str(p), "pages": len(reader.pages)}
    if meta_extra:
        meta.update(meta_extra)
    return Document(
        doc_id=str(uuid.uuid4()),
        source="pdf",
        title=(title_override or p.name),
        raw_text=text,
        meta=meta,
    )


def ingest_html_url(url: str) -> Document:
    # lightweight: fetch + strip tags. For heavier crawling, replace with Playwright/Unstructured.
    import httpx
    from bs4 import BeautifulSoup

    r = httpx.get(url, timeout=60, follow_redirects=True)
    r.raise_for_status()
    soup = BeautifulSoup(r.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    text = "\n".join([line.strip() for line in soup.get_text("\n").splitlines() if line.strip()])
    title = (soup.title.string.strip() if soup.title and soup.title.string else url)
    return Document(
        doc_id=str(uuid.uuid4()),
        source="html",
        title=title,
        raw_text=text,
        meta={"url": url},
    )


def ingest_youtube(url: str) -> Document:
    import httpx
    from urllib.parse import parse_qs
    from youtube_transcript_api import YouTubeTranscriptApi

    # extract video id
    vid = None
    u = urlparse(url)
    if u.hostname in ("www.youtube.com", "youtube.com", "m.youtube.com"):
        vid = parse_qs(u.query).get("v", [None])[0]
    elif u.hostname in ("youtu.be",):
        vid = u.path.lstrip("/")
    if not vid:
        raise ValueError("Cannot parse YouTube video id from URL")

    # transcript
    # youtube-transcript-api changed API in v1.x:
    # - older versions: YouTubeTranscriptApi.get_transcript(video_id, ...)
    # - newer versions: YouTubeTranscriptApi().fetch(video_id, ...)
    if hasattr(YouTubeTranscriptApi, 'get_transcript'):
        transcript = YouTubeTranscriptApi.get_transcript(vid, languages=['vi', 'en'])
        text = "\n".join([t.get('text', '') for t in transcript if t.get('text')])
    else:
        api = YouTubeTranscriptApi()
        fetched = api.fetch(vid, languages=['vi', 'en'])
        # FetchedTranscript supports .to_raw_data() which returns list[dict]
        if hasattr(fetched, 'to_raw_data'):
            raw = fetched.to_raw_data()
            text = "\n".join([t.get('text', '') for t in raw if t.get('text')])
        else:
            # Fallback: best-effort join on snippet.text
            text = "\n".join([getattr(t, 'text', str(t)) for t in fetched])

    # best-effort title via oEmbed (no API key)
    yt_title = None
    try:
        o = httpx.get(
            "https://www.youtube.com/oembed",
            params={"url": url, "format": "json"},
            timeout=20,
            follow_redirects=True,
        )
        if o.status_code == 200:
            yt_title = (o.json() or {}).get("title")
    except Exception:
        yt_title = None

    title = yt_title.strip() if isinstance(yt_title, str) and yt_title.strip() else f"YouTube:{vid}"

    return Document(
        doc_id=str(uuid.uuid4()),
        source="youtube",
        title=title,
        raw_text=text,
        meta={"url": url, "video_id": vid},
    )


def _safe_filename_from_url(url: str) -> str:
    try:
        p = urlparse(url)
        name = Path(p.path).name
        if name:
            return name
    except Exception:
        pass
    return "download"


def ingest_url_auto(url: str, *, data_dir: str) -> Document:
    """Ingest a URL.

    - YouTube => transcript
    - Direct links to pdf/docx/txt/md => download then parse
    - Otherwise => treat as html
    """
    lower = (url or "").lower()
    if "youtu.be/" in lower or "youtube.com/" in lower:
        return ingest_youtube(url)

    # file extensions we can download + parse
    exts = (".pdf", ".docx", ".txt", ".md")
    if any(lower.split("?")[0].endswith(ext) for ext in exts):
        import os
        import httpx

        os.makedirs(data_dir, exist_ok=True)
        base = _safe_filename_from_url(url)
        # ensure we keep extension for parser
        ext = Path(base).suffix.lower()
        if ext not in exts:
            # if URL had extension but base didn't parse, fallback
            for e in exts:
                if lower.split("?")[0].endswith(e):
                    ext = e
                    break
        tmp = os.path.join(data_dir, f"download_{uuid.uuid4().hex}{ext or ''}")

        r = httpx.get(url, timeout=120, follow_redirects=True)
        r.raise_for_status()
        with open(tmp, "wb") as f:
            f.write(r.content)

        title_override = base if base else None
        meta_extra = {"url": url, "original_name": base}

        if tmp.endswith(".pdf"):
            return ingest_pdf_path(tmp, title_override=title_override, meta_extra=meta_extra)
        if tmp.endswith(".docx"):
            return ingest_docx_path(tmp, title_override=title_override, meta_extra=meta_extra)
        return ingest_txt_path(tmp, title_override=title_override, meta_extra=meta_extra)

    # default to html
    return ingest_html_url(url)
